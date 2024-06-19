import pandas as pd
from docx import Document
from docx.shared import Pt,Inches,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

# Function to generate Challan Letter
def generate_challan(bikeno,date):
        # Create a Word document
        doc = Document()
        
    # Add a table with a single cell
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

    # Add an image to the cell
        cell.paragraphs[0].add_run().add_picture('logo.jpg',width=Inches(6))

    # Set cell alignment to center
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        

        # Add header
        paragraph = doc.add_paragraph('TRAFFIC CHALLAN LETTER')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Access the Run object
        run = paragraph.runs[0]

    # Change font size to 14 points
        run.font.size = Pt(14)

    # Change text color to red
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add date
        doc.add_paragraph(f'Date : {date}')


        # Add main content
        doc.add_paragraph("\nDear Sir/Madam,")
        doc.add_paragraph("\nSubject : CHALLAN FOR NOT WEARING HELMET")
        doc.add_paragraph("\n      We regret to inform you that you have been issued a traffic challan for the following violation.")

        # Add details of the violation
        doc.add_paragraph("\nViolation Details :")
        doc.add_paragraph("                     - Offense : Not Wearing Helmet.")
        doc.add_paragraph("                     - Bike No : {}".format(bikeno))
        doc.add_paragraph("                     - Date and Time : 12-10-2023 10:04:32.")
        doc.add_paragraph("                     - Location : Near Ambedkhar statue,Noida.")

        # Add information about the fine
        doc.add_paragraph("\nAs per traffic regulations, not wearing a helmet is a serious offense , and you are required to pay a fine for the violation.")
        doc.add_paragraph("The fine amount is Rs.1048.04, and you are requested to submit the payment at the earliest.")

        # Add closing
        doc.add_paragraph("\nPlease ensure compliance with traffic rules and regulations to avoid further penalties.")

        doc.add_paragraph("\nSincerely,")
        doc.add_paragraph("TRAFFIC POLICE OF UTTAR PRADESH.")

        # Save the Word document
        doc.save('{}.docx'.format(bikeno))
# Example usage
